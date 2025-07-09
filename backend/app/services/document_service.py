"""
文档处理服务
负责文档格式转换、内容提取等功能
"""

import os
import aiofiles
from docx import Document
from typing import Optional
import re

from app.models.schemas import DocumentType

class DocumentService:
    """文档处理服务类"""
    
    async def convert_to_markdown(self, file_path: str, doc_type: DocumentType) -> str:
        """
        将文档转换为Markdown格式
        
        Args:
            file_path: 文件路径
            doc_type: 文档类型
            
        Returns:
            转换后的Markdown内容
        """
        try:
            if doc_type == DocumentType.MARKDOWN:
                return await self._read_markdown_file(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._convert_docx_to_markdown(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._convert_txt_to_markdown(file_path)
            else:
                raise ValueError(f"不支持的文档类型: {doc_type}")
                
        except Exception as e:
            raise Exception(f"文档转换失败: {str(e)}")
    
    async def _read_markdown_file(self, file_path: str) -> str:
        """读取Markdown文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        return content
    
    async def _convert_txt_to_markdown(self, file_path: str) -> str:
        """将TXT文件转换为Markdown"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        # 简单的TXT到Markdown转换
        # 将空行分隔的段落转换为Markdown段落
        paragraphs = content.split('\n\n')
        markdown_content = '\n\n'.join(paragraph.strip() for paragraph in paragraphs if paragraph.strip())
        
        return markdown_content
    
    async def _convert_docx_to_markdown(self, file_path: str) -> str:
        """将DOCX文件转换为Markdown"""
        try:
            doc = Document(file_path)
            markdown_lines = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测标题级别（基于样式或格式）
                style_name = paragraph.style.name.lower()
                
                if 'heading' in style_name:
                    # 提取标题级别
                    level = self._extract_heading_level(style_name)
                    markdown_lines.append(f"{'#' * level} {text}")
                elif 'title' in style_name:
                    markdown_lines.append(f"# {text}")
                else:
                    # 处理普通段落
                    formatted_text = self._format_paragraph_text(paragraph)
                    markdown_lines.append(formatted_text)
                
                markdown_lines.append("")  # 添加空行
            
            # 处理表格
            for table in doc.tables:
                markdown_table = self._convert_table_to_markdown(table)
                if markdown_table:
                    markdown_lines.extend(markdown_table)
                    markdown_lines.append("")
            
            return '\n'.join(markdown_lines).strip()
            
        except Exception as e:
            raise Exception(f"DOCX转换失败: {str(e)}")
    
    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名称中提取标题级别"""
        # 查找数字
        import re
        match = re.search(r'(\d+)', style_name)
        if match:
            level = int(match.group(1))
            return min(level, 6)  # Markdown最多支持6级标题
        return 1
    
    def _format_paragraph_text(self, paragraph) -> str:
        """格式化段落文本，保留基本格式"""
        text = paragraph.text
        
        # 这里可以添加更复杂的格式转换逻辑
        # 例如：粗体、斜体、链接等
        
        return text
    
    def _convert_table_to_markdown(self, table) -> Optional[list]:
        """将Word表格转换为Markdown表格"""
        try:
            if not table.rows:
                return None
            
            markdown_table = []
            
            # 处理表头
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            markdown_table.append("| " + " | ".join(headers) + " |")
            markdown_table.append("| " + " | ".join(["---"] * len(headers)) + " |")
            
            # 处理数据行
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                # 确保单元格数量与表头一致
                while len(cells) < len(headers):
                    cells.append("")
                markdown_table.append("| " + " | ".join(cells[:len(headers)]) + " |")
            
            return markdown_table
            
        except Exception:
            return None
    
    def validate_markdown(self, content: str) -> dict:
        """验证Markdown内容的有效性"""
        issues = []
        suggestions = []
        
        # 检查标题结构
        lines = content.split('\n')
        heading_levels = []
        
        for i, line in enumerate(lines, 1):
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_levels.append((i, level))
        
        # 检查标题级别跳跃
        for i in range(1, len(heading_levels)):
            prev_level = heading_levels[i-1][1]
            curr_level = heading_levels[i][1]
            
            if curr_level > prev_level + 1:
                issues.append(f"第{heading_levels[i][0]}行: 标题级别跳跃过大")
        
        # 检查空标题
        for line_num, level in heading_levels:
            line = lines[line_num - 1]
            if len(line.strip()) <= level + 1:  # 只有#号没有内容
                issues.append(f"第{line_num}行: 空标题")
        
        # 提供优化建议
        if len(heading_levels) == 0:
            suggestions.append("建议添加标题来组织文档结构")
        
        if len([line for line in lines if line.strip()]) < 10:
            suggestions.append("文档内容较少，建议丰富内容")
        
        return {
            "valid": len(issues) == 0,
            "issues": issues,
            "suggestions": suggestions,
            "stats": {
                "lines": len(lines),
                "headings": len(heading_levels),
                "paragraphs": len([line for line in lines if line.strip() and not line.strip().startswith('#')])
            }
        }
